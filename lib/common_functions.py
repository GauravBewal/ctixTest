import datetime
import json
import random
import shutil
import sys
from time import sleep
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.firefox_profile import *
from selenium.webdriver import FirefoxOptions
from selenium.webdriver.firefox.options import Options as FOptions
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from collections import defaultdict
from config.process_config import *
import os.path
from os import path
import glob
import xml.etree.ElementTree as et
import HtmlTestRunner
import xmlrunner as xmlrunner
from sys import platform
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.events import EventFiringWebDriver, AbstractEventListener
import re
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.firefox import GeckoDriverManager
import csv
from docx import Document
from docx.shared import Inches, Pt


class SuperListener(AbstractEventListener):
    """
    The motive of this listener is to capture all kind of exceptions or actions and then implement a fail safe around
    them. This can be used to tell the code what will happen if some certain exception occurred.
    Todo: This is a very powerful custom class, we can implement screenshot capturing on other exceptions aswell
    """

    def on_exception(self, exception, driver):
        print("on_exception")
        # fprint(self, "[Failed] Exception occurred")


# Dynamic wait + Raise clean failure if element not found
def waitfor(self, timeout, attribute, value, raisefailure=True):
    """
    Waitfor method is a dynamic wait method which relies directly on the element to appear on the screen.
    It should be usually used after navigating to a new page. This provides stability to the automation scripts.
    :param self: Self Instance
    :param timeout: Number of seconds to wait until the element appears.
    :param attribute: What property to use to locate the element, eg By.NAME , By.XPATH etc
    :param value: Value of the property to be used, eg for XPATH "//div[@class,'classname']" etc
    :param raisefailure: If you want the test case to fail if the element is not found. By default this is set to True.
    :return: The method returns if the element is found or not as a boolean value of True/False
    """
    # raisefailure yes= 1
    # raisefailure no = 0
    driver = self.driver
    thinktime = 1
    status = False
    sleep(1)
    while True:

        try:
            if driver.find_elements(attribute, value).__len__() > 0:
                status = True
            else:
                status = False
        except:
            status = False

        if status:
            break
        else:
            thinktime = thinktime + 1
            sleep(1)

        if thinktime > timeout:
            if raisefailure:
                fprint(self, "TimeOut: [Failed] Element " + str(attribute) + " with value " + str(value) + " not found")
                self.fail(
                    "Case Status: [Failed] Element " + str(attribute) + " with value " + str(value) + " not found")
            # else:
            #    print("debug: Element " + attribute + " with value " + value + " not found")
            break
    sleep(1)
    return status


def fprint(self, msg, log_level=0):  # Formatted and Structured Log
    """
    fprint is an advance implementation of print method, fprint = Formatted print. This single function can be used to
    direct any kind of print messages. This is token sensitive if in any message [Failed] is passed the program will
    capture a screenshot and add into the report.
    :param self: Self Instance
    :param msg: Pass the message to print
    :param log_level: By default it is 0 - Verbose
    """

    if msg.__contains__("TC_ID"):
        testid = re.findall(r'-?\d+\.?\d*', msg)
        if len(testid) > 0:
            testid = testid[0]
        else:
            testid = "0"
        author = get_value(str(os.path.abspath(sys.argv[0])))
        if author is None:
            author = ""
        print("*** FilePath::" + str(os.path.abspath(sys.argv[0])) + " ***<br>")
        print("*** " + str(self.id()).split('.')[-1] + "::" + testid + " ::author::" + str(
            author) + "::***<br>")  # This entry will be read in email reporting

    # We can create a more advanced logger function to analyze and process logs.
    if report_format == "xml":
        print(msg)
    else:
        print(msg + "<br>")
    if str(msg).__contains__("[Failed]"):
        takescreenshot(self)
    # if log_level == 0:  # Verbose
    # Verbose
    # Debug
    # Production
    # Error
    # Exceptions
    # Lib issues
    # So many possibilities here :)


def initialize_browser(self):
    """
    This method initializes the Browser based on the settings done in config.ini.
    It supports chrome and firefox for now.
    WebDriver needs to be present in the environment/PYTHONPATH.
    :param self: Self Instance
    """
    video_timer(self, True)
    if RunOnBrowser == "firefox":
        profile = webdriver.FirefoxProfile()
        profile.set_preference('browser.download.folderList', 2)
        profile.set_preference('browser.download.manager.showWhenStarting', False)
        # profile.set_preference('browser.download.dir', os.getcwd())
        dir_path = os.path.join(os.environ["PYTHONPATH"], "reports", "downloadFiles")
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
        os.mkdir(dir_path)
        profile.set_preference('browser.download.dir', dir_path)
        # profile.set_preference('browser.download.dir', 'Untitled/Users/gauravbewal/Documents')
        profile.set_preference('browser.helperApps.neverAsk.saveToDisk', 'text/csv, application/csv')  # Mime type
        # profile.set_preference('pdfjs.disabled', True)
        '''
        op = FOptions()
        op.add_argument("--headless")
        # Findout if the execution is happening inside the docker(jenkins) or not
        if path.exists("/.dockerenv"):
            print("Execution Platform is Docker Container")
            driver_path = "/usr/local/bin/geckodriver"
        elif platform == "linux" or platform == "linux2":
            print("Execution Platform is Linux")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "linux", "geckodriver")
        elif platform == "darwin":
            print("Execution Platform is MacOS")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "macos", "geckodriver")
        elif platform == "win32":
            print("Execution Platform is Windows")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "win", "geckodriver.exe")
        else:
            print("Could not identify the platform type - Issues might happen for Chromedriver")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "macos", "geckodriver")
        self.driver = webdriver.Firefox(options=op, firefox_profile=profile, executable_path=driver_path,
                                        service_log_path="/tmp/geckodriver.log")
        
        self.driver = webdriver.Firefox(options=op, firefox_profile=profile,
                                        executable_path=GeckoDriverManager().install(),
                                        service_log_path="/tmp/geckodriver.log")
        '''
        if path.exists("/.dockerenv"):
            print("Execution Platform is Docker Container")
            driver_path = "/usr/local/bin/geckodriver"
            self.driver = webdriver.Firefox(firefox_profile=profile, executable_path=driver_path,
                                            service_log_path=os.path.join(report_location, "geckodriver.log"))
        else:
            self.driver = webdriver.Firefox(executable_path=GeckoDriverManager().install(), firefox_profile=profile)
    else:
        op = Options()
        op.add_argument("–-no-sandbox")
        # op.add_argument('--headless')
        op.add_argument("--log-level=1")
        op.add_argument("-–disable-dev-shm-usage")
        op.add_argument("--disable-extensions")
        op.add_experimental_option("prefs", {
            "download.default_directory": os.path.join(os.environ["PYTHONPATH"], "reports", "downloadFiles")})
        op.set_capability(
            "goog:loggingPrefs", {"performance": "ALL", "browser": "ALL", "client":"ALL", "server":"ALL", "driver":"ALL"}
        )
        # op.add_argument("--disable - gpu") #temporary fix chromedriver issue on 92 version
        '''
        if path.exists("/.dockerenv"):
            print("Execution Platform is Docker Container")
            driver_path = "/usr/local/bin/chromedriver"
        elif platform == "linux" or platform == "linux2":
            print("Execution Platform is Linux")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "linux", "chromedriver")
        elif platform == "darwin":
            print("Execution Platform is MacOS")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "macos", "chromedriver")
        elif platform == "win32":
            print("Execution Platform is Windows")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "win", "chromedriver.exe")
        else:
            print("Could not identify the platform type - Issues might happen for Chromedriver")
            driver_path = os.path.join(os.environ["PYTHONPATH"], "tools", "macos", "chromedriver")
        self.driver = webdriver.Chrome(driver_path, options=op)
        '''
        dc = DesiredCapabilities.CHROME
        dc['goog:loggingPrefs'] = {'browser': 'ALL'}
        if path.exists("/.dockerenv"):
            print("Execution Platform is Docker Container")
            driver_path = "/usr/local/bin/chromedriver"
            self.driver = webdriver.Chrome(driver_path, options=op, desired_capabilities=dc)
        else:
            self.driver = webdriver.Chrome(ChromeDriverManager().install(), options=op, desired_capabilities=dc)
    # This functionality is fixed in Selenium 4 python binding , beta 3 build has this fix
    # self.driver = EventFiringWebDriver(self.driver, SuperListener())
    self.driver.maximize_window()
    return self.driver


def process_console_logs(self):
    if RunOnBrowser == "chrome" and testplan != 'prodsanity.txt':
        sleep(2)
        Fail = False
        msg = ""
        fprint(self, "Checking Browser console logs for errors")
        for entry in self.driver.get_log('browser'):
            if str(entry).__contains__("SEVERE"):
                msg = msg + str(entry)
                Fail = True
        if Fail:
            fprint(self, "Failed - Issues found in Browser console logs.")
            self.fail("[Failed] - Browser Console Debug Logs: " + msg)
        else:
            fprint(self, "Passed - No issues found in Browser console logs.")
    sleep(0.5)


def clear_console_logs(self):
    """
    This function helps to clear the browser logs. There is no proper option with Chrome driver to clear the logs,
    but calling get log on chrome driver basically sets the timestamp at a current date and older logs are ignored.
    This function should be called at the beginning/setUp of the new test case or at the ending/teardown of the case.
    """
    fprint(self, "--Initializing test: Cleaning browser console logs for previous errors--")
    if RunOnBrowser == "chrome":
        self.driver.get_log('browser')
        sleep(0.5)


def takescreenshot(self, sname=None):
    """
    When this function is called it takes the screenshot/png image of the browser right away and
     save it to the configured reporting location with datetime stamp.
    """
    now = sname if sname else datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    name = os.path.join(report_location, now + ".png")
    self.driver.save_screenshot(name)
    # Temp Solution until we have a reporting portal. Running in docker/jenkins? put clickable links for Jenkins
    if report_format == "xml" and path.exists("/.dockerenv"):
        fprint(self,
               "Screenshot: https://jenkins.cyware.com/job/" + jenkins_job_name + "/HTML_20Report/" + now + ".png")
    else:
        fprint(self, "<a href='" + now + ".png" + "'>Screenshot</a><br>")


def reporting():
    """
    This function returns the Runner instance of either html reporting or xml reporting based on settings done.
    :return: HTML OR XML Runner configured at reporting location defined in config.ini
    """
    # Todo: Pass the script name for reporting portal
    # currentTime = str(datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S'))
    # scriptName = sys.argv[0].split('.')[0]
    # testResultsName = scriptName + '_' + currentTime
    if report_format == "xml":
        return xmlrunner.XMLTestRunner(output=report_location)
    elif report_format == "html":
        return HtmlTestRunner.HTMLTestRunner(output=report_location)
    elif report_format == "none":
        return None
    else:
        return xmlrunner.XMLTestRunner(output=report_location)


def set_value(key, value):
    if not os.path.exists(temp_settings_path):
        # Create settings file
        try:
            temp_file = open(temp_settings_path, "w")
            temp_file.write("{\"key\":\"value\"}")
            temp_file.close()
        except Exception as x:
            print(str(x))
    else:
        # File exist, but empty
        if os.stat(temp_settings_path).st_size == 0:
            # Create settings file
            with open(temp_settings_path, "w") as f:
                f.write("{\"key\":\"value\"}")
                f.close()

    with open(temp_settings_path, "r") as f:
        settings = eval(f.read())

    with open(temp_settings_path, "w") as f:
        settings[key] = value
        f.write(str(settings))
        f.close()


def get_value(key):
    if not os.path.exists(temp_settings_path):
        return None
    else:
        with open(temp_settings_path, "r") as f:
            settings = eval(f.read())
            f.close()

        if key in settings.keys():
            return settings[key]
        else:
            return None


def clear_field(element):
    while len(element.get_attribute("value")) > 0:
        element.send_keys(Keys.ARROW_RIGHT)
        element.send_keys(Keys.BACK_SPACE)
    sleep(1)


def search(self, value):
    waitfor(self, 20, By.XPATH, "//input[@placeholder='Search or filter results']")
    fprint(self, "Search bar is visible, searching - " + value)
    self.driver.find_element_by_xpath("//input[@placeholder='Search or filter results']").click()
    self.driver.find_element_by_xpath("//input[@placeholder='Search or filter results']").clear()
    self.driver.find_element_by_xpath("//input[@placeholder='Search or filter results']").send_keys(value)
    self.driver.find_element_by_xpath("//span[contains(text(),'Press enter or click to search')]").click()


def choose_previous_date(self):
    fprint(self, "Selecting date from the calender")
    waitfor(self, 20, By.XPATH, "//div[contains(@class,'cy-date-picker--')]")
    self.driver.find_element_by_xpath(
        "//div[contains(@class,'cy-date-picker--')]").click()
    if waitfor(self, 5, By.XPATH, "(//td[@class='available today current']/preceding::td[1])[1]", False):
        self.driver.find_element_by_xpath("(//td[@class='available today current']/preceding::td[1])[1]").click()
        fprint(self, "Selected Yesterday's date")
    elif waitfor(self, 5, By.XPATH, "//td[@class='available today current']/ancestor::tr/preceding::tr[1]/td[7]",
                 False):
        self.driver.find_element_by_xpath(
            "//td[@class='available today']/ancestor::tr/preceding::tr[1]/td[7]").click()
        fprint(self, "Selected Yesterday's date")
    else:
        fprint(self, "Going back to previous month")
        self.driver.find_element_by_xpath("//button[@aria-label='Previous Month']").click()
        waitfor(self, 5, By.XPATH, "//td[@class='next-month']")
        self.driver.find_element_by_xpath("//td[@class='next-month']").click()
        fprint(self, "Selected previous month date")


def video_timer(self, explicit=False):
    current_time = get_value("current_time")
    if current_time is None:
        set_value("current_time", str(int(time.time())))
        fprint(self, str("Video Time: 0"))
    if explicit:
        totalt = int(time.time()) - int(get_value("current_time"))
        fprint(self, str("Video Time: ") + str(time.strftime("%H:%M:%S", time.gmtime(totalt))))


def verify_success(self, message, timeout=20):
    success_message = wait_and_get_message(self, By.XPATH, "//div[contains(@class, 'cy-message__text')]", timeout,
                                           False)
    if message.lower() in success_message.lower():
        fprint(self, "[Passed] Expected message is found, " + str(message))
    elif success_message == "alert_not_found":
        fprint(self, "Case Status: [Failed] No Alert Found")
        self.fail("Case Status: [Failed]  No Alert Found")
    elif message.lower() not in success_message.lower():
        fprint(self,
               "[Failed] Alert found with different msg. Found:" + success_message + "Expected:" + message)
        self.fail("Case Status: [Failed] Alert found but expected message is not found")


def wait_and_get_message(self, attribute, value, timeout=20, raisefailure=True):
    """
     Waitfor method is a dynamic wait method which relies directly on the element to appear on the screen.
     It should be usually used after navigating to a new page. This provides stability to the automation scripts.
     :param self: Self Instance
     :param timeout: Number of seconds to wait until the element appears.
     :param attribute: What property to use to locate the element, eg By.NAME , By.XPATH etc
     :param value: Value of the property to be used, eg for XPATH "//div[@class,'classname']" etc
     :param raisefailure: If you want the test case to fail if the element is not found. By default this is set to True.
     :return: The method returns if the element is found or not as a boolean value of True/False
     """
    # raisefailure yes= 1
    # raisefailure no = 0
    driver = self.driver
    thinktime = 1
    status = False
    while True:

        try:
            if driver.find_elements(attribute, value).__len__() > 0:
                status = True
            else:
                status = False
        except:
            status = False
        if status:
            check_count = 0
            text_is = self.driver.find_element_by_xpath(value).text
            while text_is == '' and check_count < 50:
                text_is = self.driver.find_element_by_xpath(value).text
                check_count += 1
            self.driver.find_elements_by_xpath("//div[@class='el-notification__closeBtn el-icon-close']")[0].click()
            final_msg = re.sub(" +", " ", text_is)
            return final_msg
        else:
            thinktime = thinktime + 1
            sleep(1)
        if thinktime > timeout:
            if raisefailure:
                fprint(self, "TimeOut: [Failed] Element " + str(attribute) + " with value " + str(value) + " not found")
                self.fail(
                    "Case Status: [Failed] Element " + str(attribute) + " with value " + str(value) + " not found")
            break
    return "alert_not_found"


def set_credentials(cred_name, key, value):
    with open(os.path.join(os.environ["PYTHONPATH"], "creds.json"), "r") as rfile:
        data = json.load(rfile)
    if cred_name not in data.keys():
        data[cred_name] = {}
    data[cred_name][key] = value
    with open(os.path.join(os.environ["PYTHONPATH"], "creds.json"), "w") as wfile:
        wfile.write(json.dumps(data, indent=4))


def get_credentials(cred_name):
    print(f"Getting credentials for {cred_name}")
    with open(os.path.join(os.environ["PYTHONPATH"], "creds.json"), "r") as rfile:
        data = json.load(rfile)
    return data[cred_name]


def generate_random_ip():
    """
        Funtion to create random ip
    """
    val = '.'.join(str(random.randint(1, 256)) for _ in range(4))
    return val


def get_main_modules(self):
    """
        Function to fetch all modules listed under main menu
    """
    waitfor(self, 15, By.XPATH, "//i[@class='cyicon-menu']")
    self.driver.find_element_by_xpath("//i[@class='cyicon-menu']").click()
    sleep(1)
    if waitfor(self, 5, By.XPATH, "//li[@aria-haspopup='listbox']", False):
        ele = self.driver.find_elements_by_xpath("//li[@aria-haspopup='listbox']")
        [i.click() for i in ele]
    modules = self.driver.find_elements_by_xpath("//a[contains(@class,'cy-sidebar__menu-item')]")
    self.driver.find_element_by_xpath("//i[@class='cyicon-menu']").click()
    main_modules = []
    for module in modules:
        _title = module.get_attribute('aria-label')
        main_modules.append(_title) if _title else 1
    return main_modules


def get_admin_modules(self):
    """
        Function to get all modules under the Admin menu
    """
    waitfor(self, 15, By.XPATH, "//i[@class='cyicon-menu-admin']")
    self.driver.find_element_by_xpath("//i[@class='cyicon-menu-admin']").click()
    sleep(1)
    admin_modules = []
    modules = self.driver.find_elements_by_xpath("//ul[contains(@class, 'cy-sidebar__quick-menu')]/li/a")
    self.driver.find_element_by_xpath("//i[@class='cyicon-menu-admin']").click()
    for module in modules:
        _title = module.get_attribute('aria-label')
        admin_modules.append(_title) if _title else 1
    return admin_modules


def get_quick_action_modules(self):
    """
        Function to get modules listed under quick access
    """
    quick_access_path = "//ul[contains(@class,'cy-topbar__menu')]//button[normalize-space(text())='New']"
    quick_access_dropdown = "//ul[contains(@class,'cy-topbar__quick-menu__container')]//li"
    self.driver.find_element_by_xpath(quick_access_path).click()
    waitfor(self, 5, By.XPATH, quick_access_dropdown)
    modules = self.driver.find_elements_by_xpath(quick_access_dropdown)
    module_list = []
    for module in modules:
        module_list.append(module.text)
    self.driver.find_element_by_xpath(quick_access_path).click()
    return module_list


def set_snaps(module, snap_name):
    """
        Setting snap names to json file
    """
    with open(os.path.join(os.environ["PYTHONPATH"], "sanity_modules.json"), "r") as rfile:
        data = json.load(rfile)
    if module not in data.keys():
        data[module] = []
    data[module].append(snap_name)
    with open(os.path.join(os.environ["PYTHONPATH"], "sanity_modules.json"), "w") as wfile:
        wfile.write(json.dumps(data, indent=4))


def get_snaps():
    """
        Getting json data of snaps captured in sanity
    """
    print(f"Getting data of snaps captured in sanity")
    with open(os.path.join(os.environ["PYTHONPATH"], "sanity_modules.json"), "r") as rfile:
        data = json.load(rfile)
    return data


def populate_sanity_report(document, module_list):
    """
        Function to populate snaps to sanity report
    """
    snaps = get_snaps()
    for module in module_list:
        if module not in snaps.keys():
            continue
        document.add_heading(module, 2)
        for i in snaps[module]:
            status = "PASS"
            try:
                if len(i[1])+len(i[2])+len(i[3])>0:
                    status = "FAIL"
                document.add_heading(i[0].split(":")[-1]+f" - {status}", 3)
                if i[1]:
                    document.add_paragraph("Console Errors")
                    for item in i[1]:
                        console_error = document.add_paragraph().add_run(item)
                        console_error.font.size = Pt(8)
                if i[2]:
                    document.add_paragraph("Driver Errors")
                    for item in i[2]:
                        driver_error = document.add_paragraph().add_run(item)
                        driver_error.font.size = Pt(8)
                if i[3]:
                    document.add_paragraph("Network Errors")
                    for item in i[3]:
                        network_error = document.add_paragraph().add_run(f"{str(item[0])}, {str(item[1])}, {str(item[2])} ")
                        network_error.font.size = Pt(8)
                document.add_picture(os.path.join(report_location, f"{module}<{i[0]}.png"), width=Inches(5),
                                     height=Inches(2.5))
            except:
                pass


def save_screenshots(self, module, sname, slp_time=.5):
    """
        Function to capture snaps and console errors
    """
    sleep(slp_time)   # required
    takescreenshot(self, sname=f"{module}<{sname}")
    _console_list,  _driver_list, _network_list = process_and_clear_logs(self)
    with open(os.path.join(os.environ["PYTHONPATH"], "sanity_modules.json"), "r") as rfile:
        data = json.load(rfile)
    if module not in data.keys():
        data[module] = []
    data[module].append([sname, _console_list, _driver_list, _network_list]) if sname not in data[module] else fprint(self, sname+" already exists")
    with open(os.path.join(os.environ["PYTHONPATH"], "sanity_modules.json"), "w") as wfile:
        wfile.write(json.dumps(data, indent=4))


def get_func_name(module):
    """
        Function to standardize function name
    """
    _a = ''
    for i in module:
        _a = _a + i.lower() if i.isalpha() else _a
    return _a


def process_and_clear_logs(self):
    _browser_list = []
    _driver_list = []
    _network_list = []
    fprint(self, "Checking Browser console logs for errors")
    for entry in self.driver.get_log('browser'):
        if str(entry).__contains__("SEVERE"):
            _browser_list.append(str(entry))
    try:
        for entry in self.driver.get_log('driver'):
            if '"level": "error"' in str(entry) or '"level": "fatal"' in str(entry):
                _driver_list.append(str(entry))
    except:
        pass
    try:
        _network_list = get_api_failures(self)
    except:
        pass
    return _browser_list, _driver_list, _network_list


def select_from_dropdown(self, val):
    waitfor(self, 10, By.XPATH, "(//li/span/span[contains(text(), '" + val + "')])[1]")
    self.driver.find_element_by_xpath("(//li/span/span[contains(text(), '" + val + "')])[1]").click()
    fprint(self, f"[Passed]-clicked on the {val}")


def nav_tabs_readonly(self, module, tabs):
    for _tab_name in tabs:
        _ele = self.driver.find_element_by_xpath(f"//li[@role='tab']/span[text()='{_tab_name}']").click()
        waitfor(self, 2, By.XPATH, f"//div[normalize-space(text())='{_tab_name}']")
        save_screenshots(self, module=module, sname=_tab_name)
    self.driver.find_element_by_xpath('//span[@data-testaction="slider-close"]').click()
    sleep(.5)  # required
    self.driver.find_element_by_xpath("//button[span[normalize-space(text())='Intel History']]").click()
    waitfor(self, 2, By.XPATH, "//div[normalize-space(text())='Intel History']", False)
    save_screenshots(self, module=module, sname="Intel History Slider")
    self.driver.find_element_by_xpath('//span[@data-testaction="slider-close"]').click()


def get_api_failures(self):
    """
        Get timeing and status of all API called in automation
    """
    responses = []
    error_list = [400, 401, 402, 403, 404, 405, 500, 501, 502, 503, 504]
    perfLog = self.driver.get_log('performance')
    for logIndex in range(0, len(perfLog)):  # Parse the Chrome Performance logs
        logMessage = json.loads(perfLog[logIndex]["message"])["message"]
        if logMessage["method"] == "Network.responseReceived":  # Filter out HTTP responses
            response = logMessage["params"]["response"]
            if "ctixapi" in response["url"] and response["status"] in error_list:
                _l = [response["status"], response["timing"]["receiveHeadersEnd"], response["url"]]
                responses.append(_l)
    return responses


video_timer(None)
